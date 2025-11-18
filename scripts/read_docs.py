#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import sys
import os

def read_excel_file(filepath):
    """Try to read Excel file using available libraries"""
    print(f"=== Reading Excel file: {filepath} ===\n")

    # Try openpyxl
    try:
        import openpyxl
        wb = openpyxl.load_workbook(filepath)
        for sheet_name in wb.sheetnames:
            print(f"\n--- Sheet: {sheet_name} ---")
            sheet = wb[sheet_name]
            for row in sheet.iter_rows(values_only=True):
                if any(cell is not None for cell in row):
                    print(row)
        return True
    except ImportError:
        pass
    except Exception as e:
        print(f"Error with openpyxl: {e}")

    # Try pandas
    try:
        import pandas as pd
        excel_file = pd.ExcelFile(filepath)
        for sheet_name in excel_file.sheet_names:
            print(f"\n--- Sheet: {sheet_name} ---")
            df = pd.read_excel(filepath, sheet_name=sheet_name)
            print(df.to_string())
        return True
    except ImportError:
        pass
    except Exception as e:
        print(f"Error with pandas: {e}")

    print("No suitable library found for reading Excel files")
    return False

def read_docx_file(filepath):
    """Try to read Word document using available libraries"""
    print(f"\n\n=== Reading Word document: {filepath} ===\n")

    # Try python-docx
    try:
        from docx import Document
        doc = Document(filepath)

        for para in doc.paragraphs:
            if para.text.strip():
                print(para.text)

        # Print tables if any
        if doc.tables:
            print("\n--- Tables ---")
            for i, table in enumerate(doc.tables):
                print(f"\nTable {i+1}:")
                for row in table.rows:
                    cells = [cell.text for cell in row.cells]
                    print(" | ".join(cells))
        return True
    except ImportError:
        pass
    except Exception as e:
        print(f"Error with python-docx: {e}")

    print("No suitable library found for reading Word documents")
    return False

if __name__ == "__main__":
    base_path = "/home/user/SCMS2"

    excel_file = os.path.join(base_path, "중앙 2팀 요구사항정의서 (1).xlsx")
    docx_file = os.path.join(base_path, "01_학생역량관리시스템_프로젝트수행계획서_v1.0_20251027 (1).docx")

    if os.path.exists(excel_file):
        read_excel_file(excel_file)
    else:
        print(f"Excel file not found: {excel_file}")

    if os.path.exists(docx_file):
        read_docx_file(docx_file)
    else:
        print(f"Word file not found: {docx_file}")
